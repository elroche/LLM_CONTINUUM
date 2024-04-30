# Importations
from flask import Flask, render_template, request, redirect, url_for
import docx
from langchain_community.llms import Ollama
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document


# Initialisation de l'instance Ollama
ollama = Ollama(base_url='http://localhost:11434', model="llama2")
oembed = OllamaEmbeddings(base_url="http://localhost:11434", model="llama2")

# Fonction lisant les documents Word (.docx)
def read_docx(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


# Application Flask pour télécharger, analyser et répondre à des questions sur des fichiers textuels (.txt) et Word (.docx)
# Création d'une application Flask
app = Flask(__name__)

# Définition d'une route pour la page d'accueil
@app.route('/')
def index():
    return render_template('upload.html')

# Définition d'une route pour le téléchargement de fichiers
@app.route('/upload', methods=['POST'])
def upload():
    global vectorstore
    # Récupération de la liste des fichiers envoyés par le formulaire
    file_paths = request.files.getlist('files')
    documents = []
    
    # Parcours, vérification et lecture de chaque fichier téléchargé
    for file in file_paths:
        if file.filename.endswith('.txt'):
            content = file.read().decode('utf-8', 'ignore')
             # Ajout du contenu du fichier à la liste des documents
            documents.append(Document(page_content=content, metadata={"source": "local"}))
        elif file.filename.endswith('.docx'):
            content = read_docx(file)
             # Ajout du contenu du fichier à la liste des documents
            documents.append(Document(page_content=content, metadata={"source": "local"}))
        else:
            return render_template('error.html', message="Unsupported file format")
    
    # Vectorisation des documents
    text_splitter = CharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    docs = text_splitter.split_documents(documents)
    datastore_directory = './datastore'
    vectorstore = Chroma.from_documents(documents=docs, embedding=oembed, persist_directory=datastore_directory)

    # Redirige vers la route pour poser une question
    return redirect(url_for('ask_question'))

# Définition d'une route pour poser une question
@app.route('/ask_question', methods=['GET', 'POST'])
def ask_question():
    global vectorstore
    if request.method == 'POST':
        # Récupération de la question posée par l'utilisateur
        question = request.form['question']
        
        # Utilisation du modèle pour répondre à la question
        qachain = RetrievalQA.from_chain_type(ollama, retriever=vectorstore.as_retriever())
        response = qachain({"query": question})
        print("response : ", response)
    
        # Obtenir la réponse textuelle
        response_text = response["result"]
        print("response_text : ", response_text)
    
        # Renvoie le template HTML "responses.html" avec la question et la réponse
        return render_template('responses.html', question=question, response=response_text)
    
    # Renvoie le template HTML "ask_question.html" pour saisir une question
    return render_template('ask_question.html')

# Démarrage de l'application Flask si ce fichier est exécuté en tant que script principal
if __name__ == "__main__":
    app.run(host='127.0.0.1', port=5002, debug=True)



