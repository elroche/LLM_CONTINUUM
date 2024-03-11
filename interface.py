from flask import Flask, render_template, request, redirect, url_for
from litellm import completion
import docx
from langchain_community.llms import Ollama
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document


# Initialisation de l'instance Ollama
ollama = Ollama(base_url='http://localhost:11434', model="llama2")
oembed = OllamaEmbeddings(base_url="http://localhost:11434", model="llama2")


def read_docx(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


app = Flask(__name__)


@app.route('/')
def index():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload():
    global vectorstore
    file_paths = request.files.getlist('files')
    documents = []

    for file in file_paths:
        if file.filename.endswith('.txt'):
            content = file.read().decode('utf-8', 'ignore')
            documents.append(Document(page_content=content, metadata={"source": "local"}))
        elif file.filename.endswith('.docx'):
            content = read_docx(file)
            documents.append(Document(page_content=content, metadata={"source": "local"}))
        else:
            return render_template('error.html', message="Unsupported file format")
    
    # Vectorisation des documents
    text_splitter = CharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    docs = text_splitter.split_documents(documents)
    datastore_directory = './datastore'
    vectorstore = Chroma.from_documents(documents=docs, embedding=oembed, persist_directory=datastore_directory)
    #vectorstore.persist(directory=datastore_directory)

    return redirect(url_for('ask_question'))

@app.route('/ask_question', methods=['GET', 'POST'])
def ask_question():
    global vectorstore
    if request.method == 'POST':
        question = request.form['question']
        
        # Utilisation du modèle pour répondre à la question
        qachain = RetrievalQA.from_chain_type(ollama, retriever=vectorstore.as_retriever())
        response = qachain({"query": question})
    
        # Obtenir la réponse textuelle
        response_text = response["response"]
    
        return render_template('responses.html', question=question, response=response)
    
    return render_template('ask_question.html')


if __name__ == "__main__":
    app.run(host='127.0.0.1', port=5002, debug=True)
